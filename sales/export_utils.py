"""
Utilidades de exportación para reportes de la tienda (PDF, Excel, Word).

Contrato de entrada: report (dict) con claves
- title: str
- subtitle: str (opcional)
- headers: List[str]
- rows: List[List[Any]]
- totals: Dict[str, Any]
- metadata: Dict[str, Any] (puede incluir periodos, currency, etc.)

Salida: BytesIO listo para enviar como attachment.
"""
from io import BytesIO
from typing import Dict, Any, List

# PDF
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# Excel
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet
from typing import cast

# Word
# Importes de python-docx dentro de la función para evitar errores si la
# dependencia aún no está instalada al momento del análisis estático.


# ============================
# Helpers comunes
# ============================

def _normalize_text(val: Any) -> str:
    return '' if val is None else str(val)


# ============================
# PDF con ReportLab
# ============================

def exportar_reporte_pdf(report: Dict[str, Any]) -> BytesIO:
    styles = getSampleStyleSheet()
    # Estilos personalizados
    styles.add(ParagraphStyle(
        name='TituloReporte', parent=styles['Title'], fontSize=18, alignment=TA_CENTER,
        textColor=colors.HexColor('#2C3E50'), spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        name='SubtituloReporte', parent=styles['Normal'], fontSize=11, alignment=TA_CENTER,
        textColor=colors.HexColor('#34495E'), spaceAfter=10
    ))
    styles.add(ParagraphStyle(
        name='Info', parent=styles['Normal'], fontSize=9, alignment=TA_LEFT,
        textColor=colors.HexColor('#7F8C8D'), spaceAfter=6
    ))

    buffer = BytesIO()
    # Ajustar a landscape si hay muchas columnas
    headers = report.get('headers') or []
    page_size = landscape(A4) if len(headers) > 6 else A4
    doc = SimpleDocTemplate(buffer, pagesize=page_size, topMargin=0.6*inch, bottomMargin=0.6*inch)
    story: List[Any] = []

    # Título y subtítulo
    title = report.get('title') or 'Reporte'
    subtitle = report.get('subtitle') or ''
    story.append(Paragraph(_normalize_text(title), styles['TituloReporte']))
    if subtitle:
        story.append(Paragraph(_normalize_text(subtitle), styles['SubtituloReporte']))

    # Metadata (periodo, moneda)
    meta = report.get('metadata') or {}
    periodo_txt = []
    if meta.get('periodo'):
        p = meta['periodo']
        fi = p.get('fecha_inicio')
        ff = p.get('fecha_fin')
        if fi or ff:
            periodo_txt.append(f"Período: {fi or 'N/A'} - {ff or 'N/A'}")
    if meta.get('currency'):
        periodo_txt.append(f"Moneda: {meta['currency']}")
    if periodo_txt:
        story.append(Paragraph(' | '.join(periodo_txt), styles['Info']))
    story.append(Spacer(1, 0.15*inch))

    # Tabla principal
    rows = report.get('rows') or []
    if headers:
        table_data = [list(map(_normalize_text, headers))]
        # Limitar filas por página es complejo; aquí renderizamos todas (ReportLab paginará si excede)
        for r in rows:
            table_data.append([_normalize_text(c) for c in r])

        # Anchos dinámicos
        total_width = 10.5 * inch  # ancho útil en landscape A4 aprox
        col_width = max(total_width / max(len(headers), 1), 1.0 * inch)
        col_widths = [col_width] * len(headers)

        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A222E')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F7F9FA')])
        ]))
        story.append(table)
        story.append(Spacer(1, 0.2*inch))

    # Totales
    totals = report.get('totals') or {}
    if totals:
        story.append(Paragraph('Totales', styles['SubtituloReporte']))
        totals_data = [['Concepto', 'Valor']]
        for k, v in totals.items():
            totals_data.append([_normalize_text(k).replace('_', ' ').title(), _normalize_text(v)])
        t2 = Table(totals_data, colWidths=[3*inch, 3*inch])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
        ]))
        story.append(t2)

    doc.build(story)
    buffer.seek(0)
    return buffer


# ============================
# Excel con openpyxl
# ============================

def exportar_reporte_excel(report: Dict[str, Any]) -> BytesIO:
    wb = Workbook()
    ws = cast(Worksheet, wb.active)
    ws.title = 'Reporte'

    # Título
    title = report.get('title') or 'Reporte'
    ws['A1'] = title
    ws['A1'].font = Font(size=16, bold=True, color='FF2C3E50')
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(report.get('headers') or [1])))
    ws['A1'].alignment = Alignment(horizontal='center')

    # Subtítulo / info
    subtitle = report.get('subtitle') or ''
    if subtitle:
        ws['A2'] = subtitle
        ws['A2'].font = Font(size=11, italic=True, color='FF34495E')

    # Encabezados
    headers = report.get('headers') or []
    rows = report.get('rows') or []
    start_row = 4
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=col, value=_normalize_text(h))
        cell.fill = PatternFill(start_color='FF1A222E', end_color='FF1A222E', fill_type='solid')
        cell.font = Font(bold=True, color='FFFFFFFF')
        cell.alignment = Alignment(horizontal='center')
        cell.border = Border(left=Side(style='thin', color='FFBDC3C7'),
                             right=Side(style='thin', color='FFBDC3C7'),
                             top=Side(style='thin', color='FFBDC3C7'),
                             bottom=Side(style='thin', color='FFBDC3C7'))

    # Filas
    for r_index, r in enumerate(rows, start=start_row + 1):
        for c_index, val in enumerate(r, start=1):
            cell = ws.cell(row=r_index, column=c_index, value=_normalize_text(val))
            cell.border = Border(left=Side(style='thin', color='FFBDC3C7'),
                                 right=Side(style='thin', color='FFBDC3C7'),
                                 top=Side(style='thin', color='FFBDC3C7'),
                                 bottom=Side(style='thin', color='FFBDC3C7'))
            if r_index % 2 == 0:
                cell.fill = PatternFill(start_color='FFF7F9FA', end_color='FFF7F9FA', fill_type='solid')

    # Totales al final
    totals = report.get('totals') or {}
    if totals:
        totals_row = start_row + 1 + len(rows) + 1
        ws.cell(row=totals_row, column=1, value='Totales')
        ws.cell(row=totals_row, column=1).font = Font(bold=True, color='FF3498DB')
        for idx, (k, v) in enumerate(totals.items(), start=0):
            ws.cell(row=totals_row + 1 + idx, column=1, value=_normalize_text(k).replace('_', ' ').title())
            ws.cell(row=totals_row + 1 + idx, column=2, value=_normalize_text(v))

    # Ajustar anchos
    for col_idx in range(1, max(2, len(headers)) + 1):
        max_len = 0
        col_letter = get_column_letter(col_idx)
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=col_idx, max_col=col_idx):
            for cell in row:
                val = '' if cell.value is None else str(cell.value)
                if len(val) > max_len:
                    max_len = len(val)
        ws.column_dimensions[col_letter].width = min(max(12, max_len + 2), 60)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


# ============================
# Word (DOCX) con python-docx
# ============================

def exportar_reporte_docx(report: Dict[str, Any]) -> BytesIO:
    # Importes locales para evitar errores si la dependencia no está instalada aún
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    doc = Document()
    # Márgenes básicos
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Título
    title = report.get('title') or 'Reporte'
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)

    # Subtítulo
    subtitle = report.get('subtitle') or ''
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(11)
            run.italic = True

    # Metadata
    meta = report.get('metadata') or {}
    lines: List[str] = []
    if meta.get('periodo'):
        fi = meta['periodo'].get('fecha_inicio')
        ff = meta['periodo'].get('fecha_fin')
        lines.append(f"Período: {fi or 'N/A'} - {ff or 'N/A'}")
    if meta.get('currency'):
        lines.append(f"Moneda: {meta['currency']}")
    if lines:
        p2 = doc.add_paragraph(' | '.join(lines))
        for run in p2.runs:
            run.font.size = Pt(9)

    # Tabla principal
    headers = report.get('headers') or []
    rows = report.get('rows') or []
    if headers:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, htxt in enumerate(headers):
            hdr_cells[i].text = _normalize_text(htxt)
            for para in hdr_cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for r in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(r):
                row_cells[i].text = _normalize_text(val)

    # Totales
    totals = report.get('totals') or {}
    if totals:
        doc.add_paragraph()  # espacio
        doc.add_heading('Totales', level=2)
        t2 = doc.add_table(rows=1, cols=2)
        t2.style = 'Light Grid Accent 1'
        hdr = t2.rows[0].cells
        hdr[0].text = 'Concepto'
        hdr[1].text = 'Valor'
        for k, v in totals.items():
            row = t2.add_row().cells
            row[0].text = _normalize_text(k).replace('_', ' ').title()
            row[1].text = _normalize_text(v)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
